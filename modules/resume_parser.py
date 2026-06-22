"""
Resume Parser Module — Extract text and structured sections from PDF/DOCX.
"""
import re
import io
import logging
from pathlib import Path
from typing import Optional

logger = logging.getLogger(__name__)

# ─── Section Keywords ─────────────────────────────────────────────────────────
SECTION_HEADERS = {
    "contact":        ["contact", "personal information", "personal details", "about me"],
    "summary":        ["summary", "objective", "profile", "about", "overview", "career objective"],
    "education":      ["education", "academic", "qualification", "degree", "university", "college", "school"],
    "experience":     ["experience", "work experience", "employment", "job history", "professional experience", "internship"],
    "skills":         ["skills", "technical skills", "competencies", "technologies", "tools", "expertise", "proficiencies"],
    "projects":       ["projects", "personal projects", "academic projects", "portfolio", "work samples"],
    "certifications": ["certifications", "certificates", "achievements", "licenses", "accreditations"],
    "awards":         ["awards", "honors", "accomplishments", "recognition", "achievements"],
    "languages":      ["languages", "spoken languages"],
    "hobbies":        ["hobbies", "interests", "activities", "extra-curricular"],
}

EMAIL_PATTERN    = re.compile(r'[\w.+-]+@[\w-]+\.[a-zA-Z]{2,}')
PHONE_PATTERN    = re.compile(r'(?:\+?\d[\d\s\-().]{7,}\d)')
LINKEDIN_PATTERN = re.compile(r'linkedin\.com/in/[\w\-]+', re.IGNORECASE)
GITHUB_PATTERN   = re.compile(r'github\.com/[\w\-]+', re.IGNORECASE)


def parse_resume(file_bytes: bytes, filename: str) -> dict:
    """
    Main entry point. Detects file type and routes to the right parser.
    Returns a structured dict of resume data.
    """
    ext = Path(filename).suffix.lower()
    if ext == ".pdf":
        raw_text = _parse_pdf(file_bytes)
    elif ext in (".docx", ".doc"):
        raw_text = _parse_docx(file_bytes)
    else:
        raw_text = file_bytes.decode("utf-8", errors="ignore")

    structured = _extract_sections(raw_text)
    structured["raw_text"] = raw_text
    structured["contact_info"] = _extract_contact(raw_text)
    structured["word_count"] = len(raw_text.split())
    return structured


def _parse_pdf(file_bytes: bytes) -> str:
    """Extract text from PDF using pdfplumber."""
    try:
        import pdfplumber
        text_parts = []
        with pdfplumber.open(io.BytesIO(file_bytes)) as pdf:
            for page in pdf.pages:
                text = page.extract_text()
                if text:
                    text_parts.append(text)
        return "\n".join(text_parts)
    except Exception as e:
        logger.error(f"PDF parse error: {e}")
        try:
            # Fallback: PyMuPDF if available
            import fitz
            doc = fitz.open(stream=file_bytes, filetype="pdf")
            return "\n".join(page.get_text() for page in doc)
        except Exception:
            return ""


def _parse_docx(file_bytes: bytes) -> str:
    """Extract text from DOCX using python-docx."""
    try:
        from docx import Document
        doc = Document(io.BytesIO(file_bytes))
        paragraphs = [p.text for p in doc.paragraphs if p.text.strip()]
        # Also grab table cells
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    if cell.text.strip():
                        paragraphs.append(cell.text)
        return "\n".join(paragraphs)
    except Exception as e:
        logger.error(f"DOCX parse error: {e}")
        return ""


def _detect_section_header(line: str) -> Optional[str]:
    """Return section key if the line looks like a section header."""
    cleaned = line.strip().lower().rstrip(":")
    for section, keywords in SECTION_HEADERS.items():
        if any(kw == cleaned or kw in cleaned for kw in keywords):
            return section
    return None


def _extract_sections(text: str) -> dict:
    """
    Split the raw resume text into named sections.
    """
    sections = {k: [] for k in SECTION_HEADERS}
    sections["other"] = []

    current_section = "other"
    lines = text.split("\n")

    for line in lines:
        stripped = line.strip()
        if not stripped:
            continue
        detected = _detect_section_header(stripped)
        if detected and len(stripped) < 60:           # actual header line
            current_section = detected
            continue
        sections[current_section].append(stripped)

    # Convert lists to text strings
    return {k: "\n".join(v) for k, v in sections.items()}


def _extract_contact(text: str) -> dict:
    """Extract email, phone, LinkedIn, GitHub."""
    emails    = EMAIL_PATTERN.findall(text)
    phones    = PHONE_PATTERN.findall(text)
    linkedins = LINKEDIN_PATTERN.findall(text)
    githubs   = GITHUB_PATTERN.findall(text)

    # Try to get name from first non-empty line
    first_line = next((l.strip() for l in text.split("\n") if l.strip()), "")
    name = first_line if len(first_line.split()) <= 5 else ""

    return {
        "name":     name,
        "email":    emails[0]    if emails    else "",
        "phone":    phones[0].strip()    if phones    else "",
        "linkedin": linkedins[0] if linkedins else "",
        "github":   githubs[0]   if githubs   else "",
    }


def get_resume_completeness(parsed: dict) -> dict:
    """Score how complete each section is (for ATS)."""
    scores = {}
    for section in ["contact", "summary", "skills", "experience", "education", "projects", "certifications"]:
        content = parsed.get(section, "")
        scores[section] = min(100, len(content.split()) * 2) if content else 0
    return scores
